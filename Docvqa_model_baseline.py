# Installation
# sudo apt-get install tesseract-ocr libtesseract-dev
# pip install pytesseract 'git+https://github.com/facebookresearch/detectron2.git' sentencepiece python-docx

# Imports
from transformers import DonutProcessor, VisionEncoderDecoderModel
import torch
import re
import json
from PIL import Image
from transformers import AutoModelForQuestionAnswering
from transformers import LayoutLMv2Processor
import os
from transformers import LayoutLMv3Processor, LayoutLMv3ForQuestionAnswering
from docx import Document
from docx.shared import Inches
import argparse


q_map={
    'q1':"Which organization issued this given circular?",
    'q2': "What is the Address of the Issuing Authority of the given Circular?",
    'q3': "What is the Serial No./ID of the Given Circular?",
    'q4': "What is the Date of Issuance of the Circular?",
    'q5': "What is the Subject of the given Circular?",
    'q6': "Who has this circular been addressed to?",
    'q7': "To Whom has the circular been forwarded to?",
    'q8': "Who Has Forwarded This Circular?",
    'q9': "What is the Designation of the Person who Forwarded this Circular?",
    'q10': "Who has signed the Given Circular?",
    'q11': "What is the Designation of the Person who Signed this Circular?"
}

def donut(model, processor, question, image):
    pixel_values = processor(image, return_tensors="pt").pixel_values
    task_prompt = "{user_input}"
    prompt = task_prompt.replace("{user_input}", question)
    decoder_input_ids = processor.tokenizer(prompt, add_special_tokens=False, return_tensors="pt")["input_ids"]

    device = "cuda" if torch.cuda.is_available() else "cpu"
    model.to(device)

    outputs = model.generate(pixel_values.to(device),
                                  decoder_input_ids=decoder_input_ids.to(device),
                                  max_length=model.decoder.config.max_position_embeddings,
                                  early_stopping=True,
                                  pad_token_id=processor.tokenizer.pad_token_id,
                                  eos_token_id=processor.tokenizer.eos_token_id,
                                  use_cache=True,
                                  num_beams=1,
                                  bad_words_ids=[[processor.tokenizer.unk_token_id]],
                                  return_dict_in_generate=True,
                                  output_scores=True)
    seq = processor.batch_decode(outputs.sequences)[0]
    seq = seq.replace(processor.tokenizer.eos_token, "").replace(processor.tokenizer.pad_token, "")
    seq = re.sub(r"<.*?>", "", seq, count=1).strip()  # remove first task start token
    return seq


def layoutlmv2(model,processor,question,image):
    encoding = processor(image, question, return_tensors="pt", padding=True, truncation=True, max_length=512)
    for k,v in encoding.items():
        encoding[k] = v.to(model.device)

    # Perform inference
    outputs = model(**encoding)
    start_logits = outputs.start_logits
    end_logits = outputs.end_logits

    # Get the answer
    start_index = torch.argmax(start_logits, dim=1).item()
    end_index = torch.argmax(end_logits, dim=1).item()

    if start_index<end_index:
      answer = encoding["input_ids"][0][start_index : end_index + 1]
      answer = processor.decode(answer)
    else:
      answer = encoding["input_ids"][0][end_index : start_index + 1]
      answer = processor.decode(answer)

    return answer


def layoutlmv3(model, processor, question, image):
    encoding = processor(image, question, return_tensors="pt", padding=True, truncation=True, max_length=512)
    for k,v in encoding.items():
        encoding[k] = v.to(model.device)

    # Perform inference
    outputs = model(**encoding)
    start_logits = outputs.start_logits
    end_logits = outputs.end_logits

    # Get the answer
    start_index = torch.argmax(start_logits, dim=1).item()
    end_index = torch.argmax(end_logits, dim=1).item()
    
    if start_index<end_index:
      answer = encoding["input_ids"][0][start_index : end_index + 1]
      answer = processor.decode(answer)
    else:
      answer = encoding["input_ids"][0][end_index : start_index + 1]
      answer = processor.decode(answer)

    return answer


def infer(args):
    with open(args.json_path,'r') as f:
        data=json.load(f)
    # load the models
    model_l2 = AutoModelForQuestionAnswering.from_pretrained("microsoft/layoutlmv2-base-uncased")
    processor_l2 = LayoutLMv2Processor.from_pretrained("microsoft/layoutlmv2-base-uncased")
    model_l3= LayoutLMv3ForQuestionAnswering.from_pretrained("microsoft/layoutlmv3-base")
    processor_l3= LayoutLMv3Processor.from_pretrained("microsoft/layoutlmv3-base")
    model_do= VisionEncoderDecoderModel.from_pretrained("naver-clova-ix/donut-base")
    processor_do= DonutProcessor.from_pretrained("naver-clova-ix/donut-base")
    
    doc=Document()
    
    for da in data:
        img_name=da['file_name']
        doc.add_heading(img_name, 1)
        img_path=os.path.join(args.img_dir,img_name)
        doc.add_picture(img_path,width=Inches(3)).alignment=1
        image = Image.open(img_path).convert("RGB")
        annotations=da['annotations']
        for annotation in annotations:
            if annotation['type']=='textarea':
                value=annotation['value']['text'][0]
                q=q_map[annotation['to_name']]
                l2=layoutlmv2(model_l2,processor_l2,q,image)
                l3=layoutlmv3(model_l3,processor_l3,q,image)
                do=donut(model_do,processor_do,q,image)
                p=doc.add_paragraph(f'Question : {q}\n >>>Orignial Answer : {value} \n >>>LayoutLMv2 : {l2}\n >>>LayoutLMv3 : {l3} \n >>>Donut : {do} \n')
                print('Question :',q)
                print('LayoutLMv2 :',l2)
                print('LayoutLMv3 :',l3)
                print('Donut :',do)
        print()
    doc.save('infer_results.docx')

def parse_args():
    parser = argparse.ArgumentParser(description="Model INFER", formatter_class=argparse.ArgumentDefaultsHelpFormatter)

    parser.add_argument("-i", "--image_dir", type=str, default=None, help="path to the image dir")
    parser.add_argument("-o", "--data_file", type=str, default="OUTPUT", help="path to the processed annotation json data file")
    args = parser.parse_args()
    return args


if __name__ == "__main__":
    args = parse_args()
    infer(args)